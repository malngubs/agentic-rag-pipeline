"""
🧠 ADVANCED AGENTIC RAG PIPELINE - PRODUCTION READY
=====================================================
A comprehensive Retrieval-Augmented Generation system that mimics human thought processes
through multi-layered cognitive architecture, adaptive reasoning, and memory systems.

Author: Advanced AI Systems
Version: 2.1 Production - OpenAI Only with Document Management
"""

import os
import asyncio
import logging
import time
import traceback
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from enum import Enum
import hashlib
import uuid
import json

# Core ML/AI Libraries
import numpy as np
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams, PointStruct

# Document Processing
import PyPDF2
import docx
import re
from nltk.tokenize import sent_tokenize
import nltk

# Download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# OpenAI
import openai

# =============================================================================
# 🏗️ CONFIGURATION & ENUMS
# =============================================================================

class ReasoningStrategy(Enum):
    """Human-like reasoning strategies"""
    DIRECT = "direct"
    ANALYTICAL = "analytical"  
    CREATIVE = "creative"
    CONSERVATIVE = "conservative"

class ConfidenceLevel(Enum):
    """Confidence levels"""
    VERY_LOW = 0.2
    LOW = 0.4
    MEDIUM = 0.6
    HIGH = 0.8
    VERY_HIGH = 0.95

@dataclass
class RAGConfig:
    """Central configuration for RAG system"""
    # Vector Database
    vector_collection_name: str = "knowledge_base"
    embedding_model: str = "all-MiniLM-L6-v2"
    vector_size: int = 384
    qdrant_host: str = "localhost"
    qdrant_port: int = 6333
    
    # Chunking Strategy
    chunk_size: int = 1000
    chunk_overlap: int = 200
    min_chunk_size: int = 100
    
    # Reasoning Parameters
    max_context_chunks: int = 5
    reasoning_depth: int = 3
    confidence_threshold: float = 0.6
    
    # LLM Settings - OpenAI Only
    openai_model: str = "gpt-4o-mini"
    openai_api_key: Optional[str] = None
    max_tokens: int = 2048
    temperature: float = 0.7

@dataclass
class MemoryChunk:
    """Knowledge piece in long-term memory"""
    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None
    confidence_score: float = 0.0
    access_count: int = 0
    last_accessed: Optional[datetime] = None

@dataclass
class DocumentMetadata:
    """Metadata for indexed documents"""
    doc_id: str
    filename: str
    file_type: str
    file_size: int
    upload_date: datetime
    chunk_count: int
    status: str  # 'indexed', 'processing', 'failed'
    error_message: Optional[str] = None

class WorkingMemory:
    """Working memory system (Miller's Magic Number: 7±2)"""
    
    def __init__(self, capacity: int = 7):
        self.capacity = capacity
        self.active_chunks: List[MemoryChunk] = []
        self.reasoning_trace: List[Dict[str, Any]] = []
        self.current_strategy: ReasoningStrategy = ReasoningStrategy.DIRECT
        self.confidence_history: List[float] = []
    
    def add_chunk(self, chunk: MemoryChunk):
        """Add chunk with capacity management"""
        self.active_chunks.append(chunk)
        if len(self.active_chunks) > self.capacity:
            self.active_chunks.sort(key=lambda x: x.last_accessed or datetime.min)
            self.active_chunks = self.active_chunks[-self.capacity:]
    
    def get_context_summary(self) -> str:
        """Get summary of active context"""
        if not self.active_chunks:
            return "No active context"
        return " | ".join([c.content[:200] for c in self.active_chunks])
    
    def clear(self):
        """Reset working memory"""
        self.active_chunks.clear()
        self.reasoning_trace.clear()
        self.confidence_history.clear()

# =============================================================================
# 📚 DOCUMENT METADATA MANAGER
# =============================================================================

class DocumentMetadataManager:
    """Manages metadata for all indexed documents"""
    
    def __init__(self, storage_file: str = "data/document_metadata.json"):
        self.logger = logging.getLogger(__name__)
        self.storage_file = Path(storage_file)
        self.documents: Dict[str, DocumentMetadata] = {}

        # ensure storage directory exists and load any existing metadata
        self.storage_file.parent.mkdir(parents=True, exist_ok=True)
        self._load_from_disk()
    
    def add_document(self, doc_id: str, filename: str, file_type: str, 
                     file_size: int, chunk_count: int, status: str = "indexed") -> DocumentMetadata:
        """Add a new document to the metadata store"""
        metadata = DocumentMetadata(
            doc_id=doc_id,
            filename=filename,
            file_type=file_type,
            file_size=file_size,
            upload_date=datetime.now(),
            chunk_count=chunk_count,
            status=status
        )
        self.documents[doc_id] = metadata
        self._save_to_disk()
        self.logger.info(f"📝 Added document metadata: {filename} (ID: {doc_id})")
        return metadata
    
    def get_document(self, doc_id: str) -> Optional[DocumentMetadata]:
        """Get document metadata by ID"""
        return self.documents.get(doc_id)
    
    def get_all_documents(self) -> List[DocumentMetadata]:
        """Get all document metadata"""
        return list(self.documents.values())
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete document metadata"""
        if doc_id in self.documents:
            del self.documents[doc_id]
            self._save_to_disk()
            self.logger.info(f"🗑️ Deleted document metadata: {doc_id}")
            return True
        return False
    
    def update_status(self, doc_id: str, status: str, error_message: Optional[str] = None):
        """Update document status"""
        if doc_id in self.documents:
            self.documents[doc_id].status = status
            if error_message:
                self.documents[doc_id].error_message = error_message
            self._save_to_disk()
    
    def get_total_chunks(self) -> int:
        """Get total number of chunks across all documents"""
        return sum(doc.chunk_count for doc in self.documents.values())
    
    def get_documents_by_status(self, status: str) -> List[DocumentMetadata]:
        """Get documents filtered by status"""
        return [doc for doc in self.documents.values() if doc.status == status]

    def _load_from_disk(self):
        """Load document metadata from disk"""
        try:
            if self.storage_file.exists():
                with open(self.storage_file, 'r') as f:
                    data = json.load(f)
                    for doc_id, doc_data in data.items():
                        doc_data['upload_date'] = datetime.fromisoformat(doc_data['upload_date'])
                        self.documents[doc_id] = DocumentMetadata(**doc_data)
                self.logger.info(f"📂 Loaded {len(self.documents)} documents from disk")
        except Exception as e:
            self.logger.warning(f"Could not load documents from disk: {e}")

    def _save_to_disk(self):
        """Save document metadata to disk"""
        try:
            data = {}
            for doc_id, doc in self.documents.items():
                doc_dict = {
                    'doc_id': doc.doc_id,
                    'filename': doc.filename,
                    'file_type': doc.file_type,
                    'file_size': doc.file_size,
                    'upload_date': doc.upload_date.isoformat(),
                    'chunk_count': doc.chunk_count,
                    'status': doc.status,
                    'error_message': doc.error_message
                }
                data[doc_id] = doc_dict
            with open(self.storage_file, 'w') as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            self.logger.error(f"Failed to save documents to disk: {e}")

# =============================================================================
# 🔤 DOCUMENT PROCESSING
# =============================================================================

class DocumentProcessor:
    """Intelligent document processing system"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
    
    async def process_file(self, file_path: str, content: bytes = None) -> List[Dict[str, Any]]:
        """Process file into semantic chunks"""
        try:
            # Extract text
            if content is not None:
                text = await self._extract_text_from_bytes(content, file_path)
            else:
                text = await self._extract_text_from_file(file_path)
            
            if not text or len(text.strip()) < self.config.min_chunk_size:
                self.logger.warning(f"Insufficient content in {file_path}")
                return []
            
            # Create chunks
            chunks = await self._create_semantic_chunks(text, file_path)
            self.logger.info(f"Processed {file_path}: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            return []
    
    async def _extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """Extract text from raw bytes - supports PDF, DOCX, Excel, and CSV"""
        extension = Path(filename).suffix.lower()

        try:
            if extension == '.pdf':
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                return "\n\n".join([page.extract_text() for page in pdf_reader.pages])

            elif extension == '.docx':
                from io import BytesIO
                doc = docx.Document(BytesIO(content))
                return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            elif extension in ['.xlsx', '.xls']:
                # Excel files - use pandas to parse
                from io import BytesIO
                import pandas as pd

                try:
                    # Read all sheets from Excel file
                    excel_file = pd.ExcelFile(BytesIO(content))
                    all_text_parts = []

                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(excel_file, sheet_name=sheet_name)

                        # Add sheet header
                        all_text_parts.append(f"=== Sheet: {sheet_name} ===")

                        # Add column names as context
                        all_text_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")

                        # Add data summary
                        all_text_parts.append(f"Total rows: {len(df)}")

                        # Convert DataFrame to readable text format
                        # Include column statistics for numeric columns
                        numeric_cols = df.select_dtypes(include=['number']).columns
                        if len(numeric_cols) > 0:
                            all_text_parts.append("\nNumeric Column Statistics:")
                            for col in numeric_cols:
                                stats = df[col].describe()
                                all_text_parts.append(f"  {col}: min={stats['min']:.2f}, max={stats['max']:.2f}, mean={stats['mean']:.2f}")

                        # Convert rows to text (limit to first 1000 rows to avoid huge documents)
                        all_text_parts.append("\nData:")
                        max_rows = min(len(df), 1000)
                        for idx, row in df.head(max_rows).iterrows():
                            row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                            all_text_parts.append(row_text)

                        if len(df) > max_rows:
                            all_text_parts.append(f"... and {len(df) - max_rows} more rows")

                        all_text_parts.append("")  # Empty line between sheets

                    return "\n".join(all_text_parts)

                except Exception as e:
                    self.logger.error(f"Excel parsing failed: {str(e)}")
                    return f"Error parsing Excel file: {str(e)}"

            elif extension == '.csv':
                # CSV files - use pandas to parse
                from io import BytesIO, StringIO
                import pandas as pd

                try:
                    # Try to detect encoding
                    text_content = content.decode('utf-8', errors='ignore')
                    df = pd.read_csv(StringIO(text_content))

                    all_text_parts = []
                    all_text_parts.append(f"=== CSV Data ===")
                    all_text_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")
                    all_text_parts.append(f"Total rows: {len(df)}")

                    # Add numeric column statistics
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        all_text_parts.append("\nNumeric Column Statistics:")
                        for col in numeric_cols:
                            stats = df[col].describe()
                            all_text_parts.append(f"  {col}: min={stats['min']:.2f}, max={stats['max']:.2f}, mean={stats['mean']:.2f}")

                    # Convert rows to text
                    all_text_parts.append("\nData:")
                    max_rows = min(len(df), 1000)
                    for idx, row in df.head(max_rows).iterrows():
                        row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        all_text_parts.append(row_text)

                    if len(df) > max_rows:
                        all_text_parts.append(f"... and {len(df) - max_rows} more rows")

                    return "\n".join(all_text_parts)

                except Exception as e:
                    self.logger.error(f"CSV parsing failed: {str(e)}")
                    # Fallback to raw text
                    return content.decode('utf-8', errors='ignore')

            else:
                # Plain text files (.txt, .md, etc.)
                return content.decode('utf-8', errors='ignore')

        except Exception as e:
            self.logger.error(f"Extraction failed: {str(e)}")
            return content.decode('utf-8', errors='ignore')
    
    async def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from file path - supports PDF, DOCX, Excel, and CSV"""
        extension = Path(file_path).suffix.lower()

        if extension == '.pdf':
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                return "\n\n".join([p.extract_text() for p in pdf.pages])

        elif extension == '.docx':
            doc = docx.Document(file_path)
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        elif extension in ['.xlsx', '.xls']:
            # Excel files - use pandas
            import pandas as pd
            try:
                excel_file = pd.ExcelFile(file_path)
                all_text_parts = []

                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    all_text_parts.append(f"=== Sheet: {sheet_name} ===")
                    all_text_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")
                    all_text_parts.append(f"Total rows: {len(df)}")

                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        all_text_parts.append("\nNumeric Column Statistics:")
                        for col in numeric_cols:
                            stats = df[col].describe()
                            all_text_parts.append(f"  {col}: min={stats['min']:.2f}, max={stats['max']:.2f}, mean={stats['mean']:.2f}")

                    all_text_parts.append("\nData:")
                    max_rows = min(len(df), 1000)
                    for idx, row in df.head(max_rows).iterrows():
                        row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        all_text_parts.append(row_text)

                    if len(df) > max_rows:
                        all_text_parts.append(f"... and {len(df) - max_rows} more rows")
                    all_text_parts.append("")

                return "\n".join(all_text_parts)
            except Exception as e:
                self.logger.error(f"Excel parsing failed: {str(e)}")
                return f"Error parsing Excel file: {str(e)}"

        elif extension == '.csv':
            # CSV files - use pandas
            import pandas as pd
            try:
                df = pd.read_csv(file_path)
                all_text_parts = []
                all_text_parts.append(f"=== CSV Data ===")
                all_text_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")
                all_text_parts.append(f"Total rows: {len(df)}")

                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    all_text_parts.append("\nNumeric Column Statistics:")
                    for col in numeric_cols:
                        stats = df[col].describe()
                        all_text_parts.append(f"  {col}: min={stats['min']:.2f}, max={stats['max']:.2f}, mean={stats['mean']:.2f}")

                all_text_parts.append("\nData:")
                max_rows = min(len(df), 1000)
                for idx, row in df.head(max_rows).iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    all_text_parts.append(row_text)

                if len(df) > max_rows:
                    all_text_parts.append(f"... and {len(df) - max_rows} more rows")

                return "\n".join(all_text_parts)
            except Exception as e:
                self.logger.error(f"CSV parsing failed: {str(e)}")
                # Fallback to raw text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

        else:
            # Plain text files
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            # Fallback
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
    
    async def _create_semantic_chunks(self, text: str, source: str) -> List[Dict[str, Any]]:
        """Create semantically meaningful chunks"""
        text = self._clean_text(text)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        chunk_id = 0
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) > self.config.chunk_size:
                if current_chunk:
                    # Generate deterministic UUID from chunk identifier
                    chunk_string_id = f"{Path(source).stem}_chunk_{chunk_id}"
                    chunk_uuid = str(uuid.uuid5(uuid.NAMESPACE_DNS, chunk_string_id))
                    chunks.append({
                        "id": chunk_uuid,
                        "text": current_chunk.strip(),
                        "source_file": source,
                        "chunk_index": chunk_id,
                        "word_count": len(current_chunk.split())
                    })
                    chunk_id += 1
                current_chunk = paragraph
            else:
                current_chunk = f"{current_chunk}\n\n{paragraph}" if current_chunk else paragraph
        
        # Final chunk
        if current_chunk and len(current_chunk.strip()) >= self.config.min_chunk_size:
            # Generate deterministic UUID from chunk identifier
            chunk_string_id = f"{Path(source).stem}_chunk_{chunk_id}"
            chunk_uuid = str(uuid.uuid5(uuid.NAMESPACE_DNS, chunk_string_id))
            chunks.append({
                "id": chunk_uuid,
                "text": current_chunk.strip(),
                "source_file": source,
                "chunk_index": chunk_id,
                "word_count": len(current_chunk.split())
            })
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = re.sub(r'\n\[Page \d+\]\n', '\n\n', text)
        return text.strip()

# =============================================================================
# 🔍 VECTOR STORE
# =============================================================================

class VectorStoreManager:
    """Intelligent vector search system"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.client: Optional[QdrantClient] = None
        self.embedding_model: Optional[SentenceTransformer] = None
        self._initialize_embedding_model()
    
    def _initialize_embedding_model(self):
        """Initialize embedding model"""
        try:
            self.logger.info(f"Loading embedding model: {self.config.embedding_model}")
            self.embedding_model = SentenceTransformer(self.config.embedding_model)
            self.logger.info("✅ Embedding model loaded")
        except Exception as e:
            self.logger.error(f"Failed to load embedding model: {str(e)}")
            raise
    
    async def initialize(self):
        """Initialize vector database connection"""
        try:
            self.logger.info(f"Connecting to Qdrant at {self.config.qdrant_host}:{self.config.qdrant_port}")
            self.client = QdrantClient(host=self.config.qdrant_host, port=self.config.qdrant_port)
            
            # Test connection
            await asyncio.get_event_loop().run_in_executor(None, self.client.get_collections)
            await self._ensure_collection_exists()
            self.logger.info("✅ Vector store initialized")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize vector store: {str(e)}")
            raise
    
    async def _ensure_collection_exists(self):
        """Ensure collection exists"""
        try:
            collections = await asyncio.get_event_loop().run_in_executor(None, self.client.get_collections)
            collection_names = [col.name for col in collections.collections]
            
            if self.config.vector_collection_name not in collection_names:
                self.logger.info(f"Creating collection: {self.config.vector_collection_name}")
                await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: self.client.create_collection(
                        collection_name=self.config.vector_collection_name,
                        vectors_config=VectorParams(size=self.config.vector_size, distance=Distance.COSINE)
                    )
                )
                self.logger.info("✅ Collection created")
            else:
                self.logger.info(f"Collection '{self.config.vector_collection_name}' exists")
                
        except Exception as e:
            self.logger.error(f"Failed to ensure collection: {str(e)}")
            raise
    
    async def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text"""
        try:
            embedding = await asyncio.to_thread(
                self.embedding_model.encode,
                text,
                show_progress_bar=False,
                normalize_embeddings=True
            )
            return embedding.tolist()
        except Exception as e:
            self.logger.error(f"Embedding generation failed: {str(e)}")
            raise
    
    async def index_chunks(self, chunks: List[Dict[str, Any]], embeddings: List[List[float]], doc_id: str) -> Dict[str, Any]:
        """Index chunks with embeddings and document ID"""
        try:
            if len(chunks) != len(embeddings):
                return {"success": False, "error": f"Mismatch: {len(chunks)} chunks vs {len(embeddings)} embeddings"}
            
            points = []
            for chunk, embedding in zip(chunks, embeddings):
                point = PointStruct(
                    id=chunk["id"],
                    vector=embedding,
                    payload={
                        "content": chunk["text"],
                        "source_file": chunk.get("source_file", ""),
                        "chunk_index": chunk.get("chunk_index", 0),
                        "word_count": chunk.get("word_count", 0),
                        "doc_id": doc_id,  # Add document ID for tracking
                        "created_at": datetime.now().isoformat()
                    }
                )
                points.append(point)
            
            # Batch upsert
            await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.upsert(collection_name=self.config.vector_collection_name, points=points)
            )
            
            self.logger.info(f"✅ Indexed {len(chunks)} chunks for document {doc_id}")
            return {"success": True, "chunks_indexed": len(chunks)}
            
        except Exception as e:
            self.logger.error(f"Indexing failed: {str(e)}")
            return {"success": False, "error": str(e)}
    
    async def search(self, query_text: str, limit: int = 5, score_threshold: float = 0.3) -> List[Dict[str, Any]]:
        """Search vector store"""
        try:
            query_embedding = await self._generate_embedding(query_text)
            
            results = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.query_points(
                    collection_name=self.config.vector_collection_name,
                    query=query_embedding,
                    limit=limit,
                    score_threshold=score_threshold
                ).points
            )
            
            return [{
                "text": hit.payload["content"],
                "source": hit.payload.get("source_file", ""),
                "score": hit.score,
                "chunk_id": hit.id,
                "doc_id": hit.payload.get("doc_id", "")
            } for hit in results]
            
        except Exception as e:
            self.logger.error(f"Search failed: {str(e)}")
            return []
    
    async def delete_document_chunks(self, doc_id: str) -> bool:
        """Delete all chunks belonging to a document"""
        try:
            from qdrant_client.http.models import Filter, FieldCondition, MatchValue
            
            # Delete points with matching doc_id
            await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.delete(
                    collection_name=self.config.vector_collection_name,
                    points_selector=Filter(
                        must=[
                            FieldCondition(
                                key="doc_id",
                                match=MatchValue(value=doc_id)
                            )
                        ]
                    )
                )
            )
            self.logger.info(f"✅ Deleted all chunks for document {doc_id}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to delete document chunks: {str(e)}")
            return False

# =============================================================================
# 🧠 LLM MANAGER - OPENAI ONLY
# =============================================================================

class LLMManager:
    """LLM manager using OpenAI exclusively for reliable performance"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.client: Optional[openai.AsyncOpenAI] = None
        self._initialize_client()
    
    def _initialize_client(self):
        """Initialize OpenAI client"""
        try:
            # Get API key from config or environment
            api_key = self.config.openai_api_key or os.getenv("OPENAI_API_KEY")
            
            if not api_key:
                self.logger.error("❌ OpenAI API key not found! Set OPENAI_API_KEY environment variable.")
                raise ValueError("OpenAI API key required")
            
            self.client = openai.AsyncOpenAI(api_key=api_key)
            self.logger.info(f"✅ OpenAI client initialized with model: {self.config.openai_model}")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize OpenAI client: {str(e)}")
            raise
    
    @property
    def current_model(self) -> str:
        """Get current model name"""
        return self.config.openai_model
    
    async def generate_response(self, query: str, context_chunks: List[str]) -> Dict[str, Any]:
        """
        Generate response using OpenAI with optimized prompt structure.
        Implements best practices from LLM optimization article:
        - Static content first (system prompt)
        - Dynamic content last (user query)
        - Clear instruction format
        """
        try:
            # Build context string from chunks
            context = "\n\n".join(context_chunks) if context_chunks else ""
            
            # OPTIMIZED PROMPT STRUCTURE: Static content first, dynamic content last
            # This enables better token caching and improved performance
            
            # System prompt (static content - cached efficiently)
            system_prompt = """You are an intelligent AI assistant with expertise in analyzing and answering questions based on provided documents.

Your capabilities:
- Provide accurate, well-structured answers based on the given context
- Cite specific information from the context when relevant
- Acknowledge when information is not available in the context
- Maintain a helpful, professional tone

Response guidelines:
- Answer directly and concisely
- Use information from the provided context
- If context doesn't contain the answer, state this clearly
- Format responses for readability"""

            # User prompt with context and query (dynamic content - at the end)
            if context:
                user_prompt = f"""Context from documents:

{context}

---

Based on the context above, please answer this question:
{query}"""
            else:
                user_prompt = f"""Question: {query}

Note: No specific context was provided. Please provide a helpful response based on general knowledge."""
            
            # Build messages array with optimized structure
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            # Generate response with error handling
            start_time = time.time()
            
            try:
                completion = await self.client.chat.completions.create(
                    model=self.config.openai_model,
                    messages=messages,
                    temperature=self.config.temperature,
                    max_tokens=self.config.max_tokens,
                    top_p=0.9,
                    frequency_penalty=0.3,
                    presence_penalty=0.3
                )
                
                response_time = time.time() - start_time
                response_text = completion.choices[0].message.content
                
                # Calculate confidence based on finish reason and context availability
                confidence = self._calculate_confidence(
                    finish_reason=completion.choices[0].finish_reason,
                    has_context=bool(context_chunks),
                    response_length=len(response_text)
                )
                
                # ✨ CAPTURE TOKEN USAGE from OpenAI API response
                token_usage = {
                    "input_tokens": completion.usage.prompt_tokens,
                    "output_tokens": completion.usage.completion_tokens,
                    "total_tokens": completion.usage.total_tokens
                }
                
                self.logger.info(f"✅ Generated response in {response_time:.2f}s (confidence: {confidence:.2f}, tokens: {token_usage['total_tokens']})")
                
                return {
                    "success": True,
                    "response": response_text,
                    "confidence": confidence,
                    "model_used": self.config.openai_model,
                    "context_chunks_used": len(context_chunks),
                    "response_time": response_time,
                    "token_usage": token_usage  # ✨ Include token usage for analytics
                }
                
            except openai.APIError as e:
                self.logger.error(f"OpenAI API error: {str(e)}")
                return {
                    "success": False,
                    "error": f"API error: {str(e)}",
                    "response": "I'm experiencing technical difficulties connecting to the AI service. Please try again."
                }
                
            except openai.RateLimitError as e:
                self.logger.error(f"OpenAI rate limit exceeded: {str(e)}")
                return {
                    "success": False,
                    "error": "Rate limit exceeded",
                    "response": "I'm currently handling many requests. Please try again in a moment."
                }
                
            except openai.APIConnectionError as e:
                self.logger.error(f"OpenAI connection error: {str(e)}")
                return {
                    "success": False,
                    "error": f"Connection error: {str(e)}",
                    "response": "I'm having trouble connecting to the AI service. Please check your internet connection and try again."
                }
                
        except Exception as e:
            self.logger.error(f"Response generation failed: {str(e)}")
            self.logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "success": False,
                "error": str(e),
                "response": "I encountered an unexpected error. Please try again."
            }
    
    def _calculate_confidence(self, finish_reason: str, has_context: bool, response_length: int) -> float:
        """Calculate confidence score based on response characteristics"""
        base_confidence = 0.8
        
        # Adjust based on finish reason
        if finish_reason == "stop":
            confidence = base_confidence
        elif finish_reason == "length":
            confidence = base_confidence * 0.9  # Response was truncated
        else:
            confidence = base_confidence * 0.7
        
        # Boost if we have context
        if has_context:
            confidence = min(0.95, confidence + 0.1)
        
        # Adjust based on response length (very short might be less confident)
        if response_length < 50:
            confidence *= 0.9
        
        return round(confidence, 2)

# =============================================================================
# 🚀 RAG SYSTEM
# =============================================================================

class RAGSystem:
    """Advanced RAG System orchestrator with document management"""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.initialized = False
        
        # Components
        self.document_processor = DocumentProcessor(config)
        self.vector_store = VectorStoreManager(config)
        self.llm_manager = LLMManager(config)
        self.metadata_manager = DocumentMetadataManager()
        
        # Statistics
        self.documents_indexed = 0
        self.total_chunks = 0
    
    @property
    def embedding_model(self):
        """Access to embedding model"""
        return self.vector_store.embedding_model
    
    async def initialize(self) -> Dict[str, Any]:
        """Initialize all components"""
        try:
            self.logger.info("🚀 Initializing RAG System (OpenAI-powered)...")

            # Initialize vector store
            await self.vector_store.initialize()

            # Verify OpenAI client is ready
            if not self.llm_manager.client:
                raise Exception("OpenAI client not initialized")

            # Recalculate document statistics from existing metadata
            await self._recalculate_stats()

            self.initialized = True

            self.logger.info("🎉 RAG System initialized successfully!")
            return {"initialized": True, "status": "ready", "llm": "OpenAI GPT-4o-mini"}

        except Exception as e:
            self.logger.error(f"❌ Initialization failed: {str(e)}")
            self.initialized = False
            return {"initialized": False, "error": str(e)}
    
    async def process_document(self, file_path: Path, content: bytes) -> bool:
        """Process document end-to-end with metadata tracking"""
        if not self.initialized:
            self.logger.error("RAG System not initialized")
            return False
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        filename = file_path.name
        file_type = file_path.suffix.lower()
        file_size = len(content)
        
        try:
            self.logger.info(f"📋 Processing document: {filename} (ID: {doc_id})")
            
            # Step 1: Extract and chunk
            chunks = await self.document_processor.process_file(str(file_path), content)
            if not chunks:
                self.logger.error("No chunks created")
                self.metadata_manager.add_document(
                    doc_id, filename, file_type, file_size, 0, "failed"
                )
                return False
            
            self.logger.info(f"✅ Created {len(chunks)} chunks")
            
            # Step 2: Generate embeddings
            chunk_texts = [chunk["text"] for chunk in chunks]
            embeddings = await asyncio.to_thread(
                self.embedding_model.encode,
                chunk_texts,
                show_progress_bar=False,
                normalize_embeddings=True
            )
            embeddings = embeddings.tolist()
            
            self.logger.info(f"✅ Generated {len(embeddings)} embeddings")
            
            # Step 3: Index in vector store with document ID
            result = await self.vector_store.index_chunks(chunks, embeddings, doc_id)
            if not result.get("success"):
                self.logger.error(f"Indexing failed: {result.get('error')}")
                self.metadata_manager.add_document(
                    doc_id, filename, file_type, file_size, len(chunks), "failed"
                )
                return False
            
            # Step 4: Store metadata
            self.metadata_manager.add_document(
                doc_id, filename, file_type, file_size, len(chunks), "indexed"
            )
            
            # Update stats
            self.documents_indexed += 1
            self.total_chunks += len(chunks)
            
            self.logger.info(f"🎉 Document processed successfully! Total: {self.documents_indexed} docs, {self.total_chunks} chunks")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Processing failed: {str(e)}\n{traceback.format_exc()}")
            self.metadata_manager.add_document(
                doc_id, filename, file_type, file_size, 0, "failed"
            )
            self.metadata_manager.update_status(doc_id, "failed", str(e))
            return False
    
    async def query(self, query_text: str, use_rag: bool = True) -> Dict[str, Any]:
        """Query the RAG system with OpenAI"""
        try:
            self.logger.info(f"🔍 Query: {query_text}")
            
            if not self.initialized or not use_rag:
                return {
                    "response": "RAG system not available. Please upload documents first and ensure OpenAI API key is configured.",
                    "sources": [],
                    "using_rag": False,
                    "context_found": 0,
                    "search_results": []  # ← ADD THIS
                }
            
            # Step 1: Search vector store for relevant context
            search_results = await self.vector_store.search(
                query_text, 
                limit=self.config.max_context_chunks,
                score_threshold=self.config.confidence_threshold
            )
            
            if not search_results:
                self.logger.warning("No relevant context found in knowledge base")
                # Still generate response, but without context
                llm_result = await self.llm_manager.generate_response(query_text, [])
                
                return {
                    "response": llm_result.get("response", "I don't have enough information to answer that question."),
                    "sources": [],
                    "using_rag": True,
                    "context_found": 0,
                    "confidence": llm_result.get("confidence", 0.3),
                    "search_results": []  # ← ADD THIS
                }
            
            self.logger.info(f"✅ Found {len(search_results)} relevant chunks (scores: {[r['score'] for r in search_results]})")
            
            # Step 2: Generate response using OpenAI with context
            context_texts = [r["text"] for r in search_results]
            sources = [r["source"] for r in search_results]
            
            llm_result = await self.llm_manager.generate_response(query_text, context_texts)
            
            if not llm_result.get("success"):
                error_msg = llm_result.get("error", "Unknown error")
                self.logger.error(f"LLM generation failed: {error_msg}")
                return {
                    "response": llm_result.get("response", "I encountered an error processing your request."),
                    "sources": [],
                    "using_rag": True,
                    "context_found": len(search_results),
                    "error": error_msg,
                    "search_results": []  # ← ADD THIS
                }
            
            return {
                "response": llm_result["response"],
                "sources": list(set(sources)),
                "using_rag": True,
                "context_found": len(search_results),
                "confidence": llm_result.get("confidence", 0.8),
                "model_used": llm_result.get("model_used", "gpt-4o-mini"),
                "search_results": search_results  # ← ADD THIS
            }
            
        except Exception as e:
            self.logger.error(f"❌ Query failed: {str(e)}")
            self.logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "response": f"I encountered an error: {str(e)}. Please try again.",
                "sources": [],
                "using_rag": False,
                "context_found": 0,
                "error": str(e),
                "search_results": []  # ← ADD THIS
            }
    
    async def generate_follow_up_questions(self, query_text: str, response_text: str, 
                                          search_results: List[Dict]) -> List[str]:
        """
        Generate intelligent follow-up questions based on the query, response, and available context
        
        Args:
            query_text: User's original question
            response_text: AI's response
            search_results: Context chunks from vector search
            
        Returns:
            List of 3-5 follow-up questions
        """
        try:
            # Build context from search results
            context_summary = "\n".join([
                f"- {result['text'][:200]}..." 
                for result in search_results[:3]
            ]) if search_results else "No additional context available"
            
            # Prompt for follow-up generation
            follow_up_prompt = f"""Based on this conversation, suggest 3-5 relevant follow-up questions the user might want to ask.

USER'S QUESTION: {query_text}

AI'S RESPONSE: {response_text[:500]}...

AVAILABLE CONTEXT:
{context_summary}

Generate 3-5 natural, conversational follow-up questions that:
1. Dig deeper into the topic
2. Ask about related information
3. Clarify details from the response
4. Explore practical applications

Format: Return ONLY the questions, one per line, no numbering or bullets.
Example:
What are the specific requirements?
How long does this process take?
Are there any exceptions to this policy?"""

            # Generate follow-up questions using OpenAI
            response = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.llm_manager.client.chat.completions.create(
                    model=self.config.openai_model,
                    messages=[{"role": "user", "content": follow_up_prompt}],
                    max_tokens=200,
                    temperature=0.7
                )
            )
            
            # Parse response
            follow_ups_text = response.choices[0].message.content.strip()
            follow_ups = [q.strip() for q in follow_ups_text.split('\n') if q.strip()]
            
            # Limit to 5 questions
            follow_ups = follow_ups[:5]
            
            self.logger.info(f"✅ Generated {len(follow_ups)} follow-up questions")
            return follow_ups
            
        except Exception as e:
            self.logger.error(f"Follow-up generation failed: {str(e)}")
            # Return default follow-ups
            return [
                "Can you tell me more about this?",
                "What else should I know?",
                "Are there any related topics?"
            ]
    
    async def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all indexed documents with metadata"""
        documents = self.metadata_manager.get_all_documents()
        return [
            {
                "doc_id": doc.doc_id,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "upload_date": doc.upload_date.isoformat(),
                "chunk_count": doc.chunk_count,
                "status": doc.status,
                "error_message": doc.error_message
            }
            for doc in documents
        ]
    
    async def get_document_details(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get details for a specific document"""
        doc = self.metadata_manager.get_document(doc_id)
        if not doc:
            return None
        
        return {
            "doc_id": doc.doc_id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "file_size": doc.file_size,
            "upload_date": doc.upload_date.isoformat(),
            "chunk_count": doc.chunk_count,
            "status": doc.status,
            "error_message": doc.error_message
        }
    
    async def delete_document(self, doc_id: str) -> Dict[str, Any]:
        """Delete a document and all its chunks"""
        try:
            # Check if document exists
            doc = self.metadata_manager.get_document(doc_id)
            if not doc:
                return {"success": False, "error": "Document not found"}
            
            # Delete chunks from vector store
            vector_deleted = await self.vector_store.delete_document_chunks(doc_id)
            if not vector_deleted:
                self.logger.warning(f"Failed to delete vector chunks for {doc_id}")
            
            # Delete metadata
            metadata_deleted = self.metadata_manager.delete_document(doc_id)
            
            # Update stats
            if metadata_deleted:
                self.documents_indexed -= 1
                self.total_chunks -= doc.chunk_count
            
            return {
                "success": True,
                "doc_id": doc_id,
                "filename": doc.filename,
                "chunks_deleted": doc.chunk_count
            }
            
        except Exception as e:
            self.logger.error(f"Failed to delete document {doc_id}: {str(e)}")
            return {"success": False, "error": str(e)}
    
    async def _recalculate_stats(self):
        """Recalculate document and chunk statistics from metadata"""
        try:
            documents = self.metadata_manager.get_all_documents()
            self.documents_indexed = sum(1 for doc in documents if doc.status == "indexed")
            self.total_chunks = sum(doc.chunk_count for doc in documents if doc.status == "indexed")
            self.logger.info(f"📊 Stats recalculated: {self.documents_indexed} documents, {self.total_chunks} chunks")
        except Exception as e:
            self.logger.warning(f"Failed to recalculate stats: {e}")
            self.documents_indexed = 0
            self.total_chunks = 0

    async def get_system_status(self) -> Dict[str, Any]:
        """Get comprehensive system status"""
        try:
            # Check Qdrant
            vector_store_connected = False
            try:
                if self.vector_store.client:
                    await asyncio.get_event_loop().run_in_executor(None, self.vector_store.client.get_collections)
                    vector_store_connected = True
            except:
                pass
            
            # Check OpenAI (simple check - client exists)
            llm_available = self.llm_manager.client is not None
            
            return {
                "initialized": self.initialized,
                "vector_store_connected": vector_store_connected,
                "llm_available": llm_available,
                "llm_provider": "OpenAI",
                "llm_model": self.config.openai_model,
                "embedding_model_loaded": self.embedding_model is not None,
                "documents_indexed": self.documents_indexed,
                "total_chunks": self.total_chunks,
                "collection_name": self.config.vector_collection_name,
                "metadata_documents": len(self.metadata_manager.get_all_documents())
            }
            
        except Exception as e:
            self.logger.error(f"Status check failed: {str(e)}")
            return {"error": str(e)}
    
    async def close(self):
        """Cleanup resources"""
        try:
            self.logger.info("Closing RAG System...")
            self.initialized = False
            self.logger.info("✅ RAG System closed")
        except Exception as e:
            self.logger.error(f"Error during close: {str(e)}")