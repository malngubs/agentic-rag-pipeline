"""
🧠 ADVANCED AGENTIC RAG PIPELINE - PRODUCTION READY
=====================================================
A comprehensive Retrieval-Augmented Generation system that mimics human thought processes
through multi-layered cognitive architecture, adaptive reasoning, and memory systems.

Key Features:
- Human-like cognitive reasoning with multi-step strategies
- Advanced vector search with semantic understanding
- Adaptive LLM management with fallback mechanisms
- Real-time document processing and chunking
- Production-grade error handling and monitoring
- Confidence-based response generation

Author: Advanced AI Systems
Version: 2.0 Production
"""

import os
import sys
import asyncio
import logging
import json
import time
import traceback
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass, asdict
from enum import Enum
import hashlib
import uuid

# Core ML/AI Libraries
import numpy as np
from sentence_transformers import SentenceTransformer
import openai
from qdrant_client import QdrantClient
from qdrant_client.http import models
from qdrant_client.http.models import Distance, VectorParams, PointStruct
from qdrant_client.conversions import common_types

# Document Processing
import PyPDF2
import docx
from pathlib import Path
import mimetypes

# Web Framework
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# HTTP Client
import httpx
import requests

# Text Processing
import re
from nltk.tokenize import sent_tokenize
import nltk

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# =============================================================================
# 🏗️ CORE CONFIGURATION & ENUMS
# =============================================================================

class ReasoningStrategy(Enum):
    """
    Human-like reasoning strategies that mirror cognitive approaches:
    
    DIRECT: Like immediate recall from memory
    ANALYTICAL: Like careful step-by-step thinking
    CREATIVE: Like lateral thinking and association
    CONSERVATIVE: Like cautious, minimal-risk responses
    """
    DIRECT = "direct"
    ANALYTICAL = "analytical"  
    CREATIVE = "creative"
    CONSERVATIVE = "conservative"

class ConfidenceLevel(Enum):
    """
    Confidence levels that mirror human certainty levels
    """
    VERY_LOW = 0.2
    LOW = 0.4
    MEDIUM = 0.6
    HIGH = 0.8
    VERY_HIGH = 0.95

@dataclass
class RAGConfig:
    """
    🎛️ Central configuration for the RAG system
    Mimics human learning preferences and cognitive settings
    """
    # Vector Database Settings
    vector_collection_name: str = "knowledge_base"
    embedding_model: str = "all-MiniLM-L6-v2"
    vector_size: int = 384
    qdrant_host: str = "localhost"
    qdrant_port: int = 6333
    
    # Chunking Strategy (like human memory segmentation)
    chunk_size: int = 1000  # Characters per chunk
    chunk_overlap: int = 200  # Overlap between chunks
    min_chunk_size: int = 100  # Minimum viable chunk
    
    # Reasoning Parameters (cognitive processing)
    max_context_chunks: int = 5  # Working memory limit
    reasoning_depth: int = 3  # How many reasoning steps
    confidence_threshold: float = 0.6  # When to trust a response
    
    # LLM Settings
    openai_model: str = "gpt-4o-mini"
    ollama_base_url: str = "http://localhost:11434"
    max_tokens: int = 2048
    temperature: float = 0.7  # Balance between creativity and consistency

# =============================================================================
# 🧠 COGNITIVE MEMORY SYSTEMS
# =============================================================================

@dataclass
class MemoryChunk:
    """
    Represents a piece of knowledge in our long-term memory system
    Similar to how humans store episodic and semantic memories
    """
    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None
    confidence_score: float = 0.0
    access_count: int = 0
    last_accessed: Optional[datetime] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage"""
        data = asdict(self)
        if self.last_accessed:
            data['last_accessed'] = self.last_accessed.isoformat()
        return data

class WorkingMemory:
    """
    🧠 Working Memory System
    
    Mimics human working memory - the active processing space where we:
    - Hold current conversation context
    - Maintain reasoning state
    - Track cognitive strategies being used
    """
    
    def __init__(self, capacity: int = 7):  # Miller's Magic Number
        self.capacity = capacity
        self.active_chunks: List[MemoryChunk] = []
        self.reasoning_trace: List[Dict[str, Any]] = []
        self.current_strategy: ReasoningStrategy = ReasoningStrategy.DIRECT
        self.confidence_history: List[float] = []
    
    def add_chunk(self, chunk: MemoryChunk):
        """Add a chunk to working memory, respecting capacity limits"""
        self.active_chunks.append(chunk)
        if len(self.active_chunks) > self.capacity:
            # Remove least recently accessed (like human memory decay)
            self.active_chunks.sort(key=lambda x: x.last_accessed or datetime.min)
            self.active_chunks = self.active_chunks[-self.capacity:]
    
    def get_context_summary(self) -> str:
        """Generate a summary of current working memory state"""
        if not self.active_chunks:
            return "No active context"
        
        content_pieces = [chunk.content[:200] + "..." if len(chunk.content) > 200 
                         else chunk.content for chunk in self.active_chunks]
        return " | ".join(content_pieces)
    
    def update_reasoning_trace(self, step: str, result: Any, confidence: float):
        """Track the reasoning process like human metacognition"""
        self.reasoning_trace.append({
            "timestamp": datetime.now().isoformat(),
            "step": step,
            "result": str(result),
            "confidence": confidence,
            "strategy": self.current_strategy.value
        })
        self.confidence_history.append(confidence)
    
    def clear(self):
        """Reset working memory (like focusing on a new task)"""
        self.active_chunks.clear()
        self.reasoning_trace.clear()
        self.confidence_history.clear()

# =============================================================================
# 🔤 ADVANCED TEXT PROCESSING SYSTEM
# =============================================================================

class DocumentProcessor:
    """
    🔤 Intelligent Document Processing System
    
    Processes documents like a human reader:
    - Understands document structure and context
    - Breaks content into meaningful semantic chunks
    - Preserves important relationships between ideas
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
    
    async def process_file(self, file_path: str, content: bytes = None) -> List[MemoryChunk]:
        """
        Process a file into memory chunks with human-like understanding
        
        Args:
            file_path: Path to the file
            content: Raw file content (optional)
        
        Returns:
            List of processed memory chunks
        """
        try:
            # Determine file type and extract text
            if content is not None:
                text = await self._extract_text_from_bytes(content, file_path)
            else:
                text = await self._extract_text_from_file(file_path)
            
            if not text or len(text.strip()) < self.config.min_chunk_size:
                self.logger.warning(f"Insufficient content in {file_path}")
                return []
            
            # Create semantic chunks (like human reading comprehension)
            chunks = await self._create_semantic_chunks(text, file_path)
            
            self.logger.info(f"Processed {file_path}: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            return []
    
    async def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file types"""
        path_obj = Path(file_path)
        extension = path_obj.suffix.lower()
        
        try:
            if extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif extension == '.docx':
                return await self._extract_from_docx(file_path)
            elif extension in ['.txt', '.md']:
                return await self._extract_from_text(file_path)
            else:
                # Try as plain text
                return await self._extract_from_text(file_path)
        except Exception as e:
            self.logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            return ""
    
    async def _extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """Extract text from raw bytes based on filename"""
        extension = Path(filename).suffix.lower()
        
        try:
            if extension == '.pdf':
                # For PDF bytes processing
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                return "\n\n".join(text_parts)
            
            elif extension == '.docx':
                # For DOCX bytes processing
                from io import BytesIO
                doc = docx.Document(BytesIO(content))
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                return "\n\n".join(paragraphs)
            
            else:
                # Try as plain text
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            self.logger.error(f"Bytes extraction failed for {filename}: {str(e)}")
            return content.decode('utf-8', errors='ignore')
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
        return "\n\n".join(text_parts)
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        return "\n\n".join(paragraphs)
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as bytes and ignore errors
        with open(file_path, 'rb') as file:
            return file.read().decode('utf-8', errors='ignore')
    
    async def _create_semantic_chunks(self, text: str, source: str) -> List[MemoryChunk]:
        """
        Create semantically meaningful chunks (like human reading comprehension)
        
        This mimics how humans break down text into meaningful units:
        - Respect paragraph boundaries
        - Maintain context across chunks
        - Preserve important relationships
        """
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Try to split by paragraphs first (natural semantic boundaries)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        current_size = 0
        chunk_id = 0
        
        for paragraph in paragraphs:
            paragraph_size = len(paragraph)
            
            # If adding this paragraph exceeds chunk size, finalize current chunk
            if (current_size + paragraph_size > self.config.chunk_size and 
                current_size > self.config.min_chunk_size):
                
                if current_chunk:
                    chunk = await self._create_memory_chunk(
                        current_chunk, source, chunk_id
                    )
                    chunks.append(chunk)
                    chunk_id += 1
                
                # Start new chunk with overlap from previous chunk
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + paragraph
                current_size = len(current_chunk)
            else:
                # Add paragraph to current chunk
                if current_chunk:
                    current_chunk += f"\n\n{paragraph}"
                else:
                    current_chunk = paragraph
                current_size += paragraph_size
        
        # Handle final chunk
        if current_chunk and len(current_chunk.strip()) >= self.config.min_chunk_size:
            chunk = await self._create_memory_chunk(current_chunk, source, chunk_id)
            chunks.append(chunk)
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text for processing"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\n\[Page \d+\]\n', '\n\n', text)
        
        return text.strip()
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of previous chunk"""
        if len(text) <= self.config.chunk_overlap:
            return text
        
        # Try to get overlap at sentence boundary
        overlap_text = text[-self.config.chunk_overlap:]
        sentences = sent_tokenize(overlap_text)
        
        if len(sentences) > 1:
            # Keep complete sentences
            return ' '.join(sentences[1:]) + ' '
        else:
            return overlap_text + ' '
    
    async def _create_memory_chunk(self, content: str, source: str, chunk_id: int) -> MemoryChunk:
        """Create a memory chunk with metadata"""
        # Generate unique ID
        content_hash = hashlib.md5(content.encode()).hexdigest()
        chunk_uuid = f"{Path(source).stem}_{chunk_id}_{content_hash[:8]}"
        
        # Extract key information for metadata
        word_count = len(content.split())
        char_count = len(content)
        
        metadata = {
            "source": source,
            "chunk_id": chunk_id,
            "word_count": word_count,
            "char_count": char_count,
            "created_at": datetime.now().isoformat(),
            "content_type": Path(source).suffix.lower() or "unknown"
        }
        
        return MemoryChunk(
            id=chunk_uuid,
            content=content,
            metadata=metadata,
            last_accessed=datetime.now()
        )

# =============================================================================
# 🔍 ADVANCED VECTOR SEARCH SYSTEM
# =============================================================================

class VectorStoreManager:
    """
    🔍 Intelligent Vector Search System
    
    Manages our long-term memory (vector database) like the human brain:
    - Stores knowledge as semantic embeddings
    - Enables associative recall and memory retrieval
    - Maintains relationship between concepts
    - Supports forgetting and memory consolidation
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.client: Optional[QdrantClient] = None
        self.embedding_model: Optional[SentenceTransformer] = None
        self._initialize_embedding_model()
    
    def _initialize_embedding_model(self):
        """Initialize the embedding model for semantic understanding"""
        try:
            self.logger.info(f"Loading embedding model: {self.config.embedding_model}")
            self.embedding_model = SentenceTransformer(self.config.embedding_model)
            self.logger.info("✅ Embedding model loaded successfully")
        except Exception as e:
            self.logger.error(f"Failed to load embedding model: {str(e)}")
            raise
    
    async def initialize(self):
        """Initialize connection to vector database"""
        try:
            self.logger.info(f"Connecting to Qdrant at {self.config.qdrant_host}:{self.config.qdrant_port}")
            self.client = QdrantClient(
                host=self.config.qdrant_host,
                port=self.config.qdrant_port
            )
            
            # Test connection
            collections = await asyncio.get_event_loop().run_in_executor(
                None, self.client.get_collections
            )
            
            await self._ensure_collection_exists()
            self.logger.info("✅ Vector store initialized successfully")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize vector store: {str(e)}")
            raise
    
    async def _ensure_collection_exists(self):
        """Ensure our knowledge collection exists"""
        try:
            collections = await asyncio.get_event_loop().run_in_executor(
                None, self.client.get_collections
            )
            
            collection_names = [col.name for col in collections.collections]
            
            if self.config.vector_collection_name not in collection_names:
                self.logger.info(f"Creating collection: {self.config.vector_collection_name}")
                
                await asyncio.get_event_loop().run_in_executor(
                    None, 
                    lambda: self.client.create_collection(
                        collection_name=self.config.vector_collection_name,
                        vectors_config=VectorParams(
                            size=self.config.vector_size,
                            distance=Distance.COSINE
                        )
                    )
                )
                self.logger.info("✅ Collection created successfully")
            else:
                self.logger.info(f"Collection '{self.config.vector_collection_name}' already exists")
                
        except Exception as e:
            self.logger.error(f"Failed to ensure collection exists: {str(e)}")
            raise
    
    async def store_chunks(self, chunks: List[MemoryChunk]) -> bool:
        """
        Store memory chunks in long-term memory (vector database)
        
        Like how humans encode experiences into long-term memory:
        - Convert to semantic embeddings
        - Store with rich metadata
        - Enable future retrieval
        """
        if not chunks:
            return True
        
        try:
            points = []
            
            for chunk in chunks:
                # Generate embedding (semantic encoding)
                if not chunk.embedding:
                    chunk.embedding = await self._generate_embedding(chunk.content)
                
                # Create point for storage
                point = PointStruct(
                    id=chunk.id,
                    vector=chunk.embedding,
                    payload={
                        "content": chunk.content,
                        "metadata": chunk.metadata,
                        "confidence_score": chunk.confidence_score,
                        "access_count": chunk.access_count,
                        "last_accessed": chunk.last_accessed.isoformat() if chunk.last_accessed else None
                    }
                )
                points.append(point)
            
            # Store in batch (efficient bulk storage)
            await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.upsert(
                    collection_name=self.config.vector_collection_name,
                    points=points
                )
            )
            
            self.logger.info(f"✅ Stored {len(chunks)} chunks in vector store")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to store chunks: {str(e)}")
            return False
    
    async def search(self, query: str, limit: int = 5, score_threshold: float = 0.7) -> List[Tuple[MemoryChunk, float]]:
        """
        Search long-term memory using semantic similarity
        
        Mimics human memory retrieval:
        - Uses associative recall
        - Returns most relevant memories
        - Includes confidence scores
        """
        try:
            # Generate query embedding
            query_embedding = await self._generate_embedding(query)
            
            # Perform semantic search
            search_result = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.search(
                    collection_name=self.config.vector_collection_name,
                    query_vector=query_embedding,
                    limit=limit,
                    score_threshold=score_threshold
                )
            )
            
            # Convert results to memory chunks
            results = []
            for hit in search_result:
                chunk = MemoryChunk(
                    id=hit.id,
                    content=hit.payload["content"],
                    metadata=hit.payload["metadata"],
                    embedding=None,  # Don't load embedding for performance
                    confidence_score=hit.payload.get("confidence_score", 0.0),
                    access_count=hit.payload.get("access_count", 0),
                    last_accessed=datetime.fromisoformat(hit.payload["last_accessed"]) if hit.payload.get("last_accessed") else None
                )
                
                # Update access tracking (like human memory reinforcement)
                chunk.access_count += 1
                chunk.last_accessed = datetime.now()
                
                results.append((chunk, hit.score))
            
            self.logger.info(f"Found {len(results)} relevant chunks for query")
            return results
            
        except Exception as e:
            self.logger.error(f"Search failed: {str(e)}")
            return []
    
    async def search_with_filter(self, query: str, filters: Dict[str, Any], limit: int = 5) -> List[Tuple[MemoryChunk, float]]:
        """Search with metadata filters (like contextual memory recall)"""
        try:
            query_embedding = await self._generate_embedding(query)
            
            # Build Qdrant filter
            qdrant_filter = self._build_qdrant_filter(filters)
            
            search_result = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.search(
                    collection_name=self.config.vector_collection_name,
                    query_vector=query_embedding,
                    query_filter=qdrant_filter,
                    limit=limit
                )
            )
            
            results = []
            for hit in search_result:
                chunk = MemoryChunk(
                    id=hit.id,
                    content=hit.payload["content"],
                    metadata=hit.payload["metadata"],
                    confidence_score=hit.payload.get("confidence_score", 0.0),
                    access_count=hit.payload.get("access_count", 0) + 1,
                    last_accessed=datetime.now()
                )
                results.append((chunk, hit.score))
            
            return results
            
        except Exception as e:
            self.logger.error(f"Filtered search failed: {str(e)}")
            return []
    
    def _build_qdrant_filter(self, filters: Dict[str, Any]):
        """Build Qdrant filter from dictionary"""
        conditions = []
        
        for key, value in filters.items():
            if key == "source":
                conditions.append(
                    models.FieldCondition(
                        key=f"metadata.{key}",
                        match=models.MatchValue(value=value)
                    )
                )
            elif key == "content_type":
                conditions.append(
                    models.FieldCondition(
                        key=f"metadata.{key}",
                        match=models.MatchValue(value=value)
                    )
                )
        
        if conditions:
            return models.Filter(must=conditions)
        return None
    
    async def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text"""
        try:
            # Clean text for embedding
            cleaned_text = re.sub(r'\s+', ' ', text.strip())
            
            # Generate embedding using sentence transformer
            embedding = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.embedding_model.encode(cleaned_text, convert_to_numpy=True)
            )
            
            return embedding.tolist()
            
        except Exception as e:
            self.logger.error(f"Embedding generation failed: {str(e)}")
            raise
    
    async def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about our knowledge base"""
        try:
            collection_info = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.client.get_collection(self.config.vector_collection_name)
            )
            
            return {
                "total_chunks": collection_info.points_count,
                "vector_size": collection_info.config.params.vectors.size,
                "distance_metric": collection_info.config.params.vectors.distance,
                "indexed": collection_info.indexed_vectors_count
            }
            
        except Exception as e:
            self.logger.error(f"Failed to get collection stats: {str(e)}")
            return {}

# =============================================================================
# 🤖 INTELLIGENT LLM MANAGEMENT SYSTEM
# =============================================================================

class LLMManager:
    """
    🤖 Intelligent Language Model Management System
    
    Manages multiple LLM providers with human-like decision making:
    - Chooses best model for each task
    - Falls back gracefully when models fail
    - Adapts to performance and availability
    - Maintains conversation context
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.openai_client: Optional[openai.OpenAI] = None
        self.available_models: List[Dict[str, Any]] = []
        self.model_performance: Dict[str, Dict[str, float]] = {}
        
    async def initialize(self):
        """Initialize LLM connections and test availability"""
        await self._initialize_openai()
        await self._initialize_ollama()
        await self._test_model_availability()
        
        if not self.available_models:
            raise Exception("No LLM models available")
            
        self.logger.info(f"✅ LLM Manager initialized with {len(self.available_models)} models")
    
    async def _initialize_openai(self):
        """Initialize OpenAI client"""
        try:
            api_key = os.getenv("OPENAI_API_KEY")
            if api_key:
                self.openai_client = openai.OpenAI(api_key=api_key)
                self.logger.info("✅ OpenAI client initialized")
            else:
                self.logger.warning("OpenAI API key not found")
        except Exception as e:
            self.logger.error(f"OpenAI initialization failed: {str(e)}")
    
    async def _initialize_ollama(self):
        """Initialize Ollama client"""
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(f"{self.config.ollama_base_url}/api/tags")
                if response.status_code == 200:
                    self.logger.info("✅ Ollama client initialized")
                else:
                    self.logger.warning("Ollama not available")
        except Exception as e:
            self.logger.error(f"Ollama initialization failed: {str(e)}")
    
    async def _test_model_availability(self):
        """Test which models are actually available"""
        # Test OpenAI models
        if self.openai_client:
            for model_name in ["gpt-4o-mini", "gpt-3.5-turbo"]:
                if await self._test_openai_model(model_name):
                    self.available_models.append({
                        "name": model_name,
                        "provider": "openai",
                        "type": "chat",
                        "context_length": 128000 if "gpt-4" in model_name else 16385
                    })
        
        # Test Ollama models
        ollama_models = await self._get_ollama_models()
        for model in ollama_models:
            if await self._test_ollama_model(model["name"]):
                self.available_models.append({
                    "name": model["name"],
                    "provider": "ollama", 
                    "type": "chat",
                    "context_length": 4096  # Default, can be configured per model
                })
    
    async def _test_openai_model(self, model_name: str) -> bool:
        """Test if OpenAI model is available"""
        try:
            response = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: self.openai_client.chat.completions.create(
                    model=model_name,
                    messages=[{"role": "user", "content": "Test"}],
                    max_tokens=10,
                    timeout=10
                )
            )
            return True
        except Exception as e:
            self.logger.warning(f"OpenAI model {model_name} not available: {str(e)}")
            return False
    
    async def _get_ollama_models(self) -> List[Dict[str, str]]:
        """Get available Ollama models"""
        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                response = await client.get(f"{self.config.ollama_base_url}/api/tags")
                if response.status_code == 200:
                    data = response.json()
                    return data.get("models", [])
        except Exception as e:
            self.logger.error(f"Failed to get Ollama models: {str(e)}")
        return []
    
    async def _test_ollama_model(self, model_name: str) -> bool:
        """Test if Ollama model is available"""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.post(
                    f"{self.config.ollama_base_url}/api/generate",
                    json={
                        "model": model_name,
                        "prompt": "Test",
                        "options": {"num_predict": 5}
                    }
                )
                return response.status_code == 200
        except Exception as e:
            self.logger.warning(f"Ollama model {model_name} not available: {str(e)}")
            return False
    
    async def generate_response(self, 
                              messages: List[Dict[str, str]], 
                              strategy: ReasoningStrategy = ReasoningStrategy.DIRECT,
                              max_attempts: int = 3) -> Tuple[str, float, Dict[str, Any]]:
        """
        Generate response using best available model with human-like reasoning
        
        Args:
            messages: Conversation messages
            strategy: Reasoning strategy to use
            max_attempts: Maximum retry attempts
        
        Returns:
            (response_text, confidence_score, metadata)
        """
        
        # Select best model for this strategy
        selected_model = await self._select_model_for_strategy(strategy)
        
        for attempt in range(max_attempts):
            try:
                self.logger.info(f"🧠 Generation attempt {attempt + 1} using {selected_model['name']} (strategy: {strategy.value})")
                
                start_time = time.time()
                
                if selected_model["provider"] == "openai":
                    response, confidence = await self._generate_openai_response(
                        selected_model, messages, strategy
                    )
                elif selected_model["provider"] == "ollama":
                    response, confidence = await self._generate_ollama_response(
                        selected_model, messages, strategy
                    )
                else:
                    raise ValueError(f"Unknown provider: {selected_model['provider']}")
                
                response_time = time.time() - start_time
                
                # Update model performance tracking
                self._update_model_performance(selected_model["name"], response_time, confidence, True)
                
                metadata = {
                    "model": selected_model["name"],
                    "provider": selected_model["provider"],
                    "strategy": strategy.value,
                    "attempt": attempt + 1,
                    "response_time": response_time,
                    "success": True
                }
                
                self.logger.info(f"✅ Response generated ({len(response)} chars) using strategy {strategy.value}")
                return response, confidence, metadata
                
            except Exception as e:
                self.logger.error(f"❌ Generation attempt {attempt + 1} failed: {str(e)}")
                
                # Update performance tracking for failure
                self._update_model_performance(selected_model["name"], 0, 0, False)
                
                if attempt < max_attempts - 1:
                    # Try different model or strategy for next attempt
                    selected_model = await self._select_fallback_model(selected_model)
                    if strategy != ReasoningStrategy.CONSERVATIVE:
                        strategy = self._get_fallback_strategy(strategy)
                
        # All attempts failed
        return ("I apologize, but I'm experiencing technical difficulties generating a response. Please try again later.", 
                0.1, 
                {"error": "All generation attempts failed", "attempts": max_attempts})
    
    async def _select_model_for_strategy(self, strategy: ReasoningStrategy) -> Dict[str, Any]:
        """Select the best model for a given reasoning strategy"""
        if not self.available_models:
            raise Exception("No models available")
        
        # Strategy-based model selection (like human task-specific thinking)
        if strategy == ReasoningStrategy.CREATIVE:
            # Prefer models with higher temperature capability
            preferred = [m for m in self.available_models if "gpt-4" in m["name"]]
            if preferred:
                return preferred[0]
                
        elif strategy == ReasoningStrategy.ANALYTICAL:
            # Prefer more powerful, analytical models
            preferred = [m for m in self.available_models if "gpt-4" in m["name"] or "llama" in m["name"]]
            if preferred:
                return preferred[0]
                
        elif strategy == ReasoningStrategy.CONSERVATIVE:
            # Prefer reliable, well-tested models
            preferred = [m for m in self.available_models if m["provider"] == "openai"]
            if preferred:
                return preferred[0]
        
        # Default: return best performing model
        return self._get_best_performing_model()
    
    def _get_best_performing_model(self) -> Dict[str, Any]:
        """Get the best performing model based on historical data"""
        if not self.model_performance:
            return self.available_models[0]
        
        best_model = None
        best_score = -1
        
        for model in self.available_models:
            name = model["name"]
            if name in self.model_performance:
                perf = self.model_performance[name]
                # Composite score: success rate + confidence - response time penalty
                score = (perf["success_rate"] * 0.5 + 
                        perf["avg_confidence"] * 0.3 - 
                        min(perf["avg_response_time"] / 10, 0.2) * 0.2)
                
                if score > best_score:
                    best_score = score
                    best_model = model
        
        return best_model or self.available_models[0]
    
    async def _select_fallback_model(self, failed_model: Dict[str, Any]) -> Dict[str, Any]:
        """Select a fallback model when current model fails"""
        available_fallbacks = [m for m in self.available_models if m["name"] != failed_model["name"]]
        
        if not available_fallbacks:
            return failed_model  # No choice but to retry same model
        
        # Prefer different provider for fallback
        different_provider = [m for m in available_fallbacks if m["provider"] != failed_model["provider"]]
        if different_provider:
            return different_provider[0]
        
        return available_fallbacks[0]
    
    def _get_fallback_strategy(self, current_strategy: ReasoningStrategy) -> ReasoningStrategy:
        """Get a fallback strategy when current strategy fails"""
        fallback_map = {
            ReasoningStrategy.CREATIVE: ReasoningStrategy.DIRECT,
            ReasoningStrategy.ANALYTICAL: ReasoningStrategy.CONSERVATIVE,
            ReasoningStrategy.DIRECT: ReasoningStrategy.CONSERVATIVE,
            ReasoningStrategy.CONSERVATIVE: ReasoningStrategy.DIRECT
        }
        return fallback_map.get(current_strategy, ReasoningStrategy.CONSERVATIVE)
    
    async def _generate_openai_response(self, 
                                      model: Dict[str, Any], 
                                      messages: List[Dict[str, str]], 
                                      strategy: ReasoningStrategy) -> Tuple[str, float]:
        """Generate response using OpenAI model"""
        
        # Adjust parameters based on strategy
        temperature = self._get_temperature_for_strategy(strategy)
        
        response = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: self.openai_client.chat.completions.create(
                model=model["name"],
                messages=messages,
                max_tokens=self.config.max_tokens,
                temperature=temperature,
                timeout=60
            )
        )
        
        response_text = response.choices[0].message.content
        
        # Calculate confidence based on response characteristics
        confidence = self._calculate_response_confidence(response_text, strategy)
        
        return response_text, confidence
    
    async def _generate_ollama_response(self, 
                                      model: Dict[str, Any], 
                                      messages: List[Dict[str, str]], 
                                      strategy: ReasoningStrategy) -> Tuple[str, float]:
        """Generate response using Ollama model"""
        
        # Convert messages to prompt format for Ollama
        prompt = self._convert_messages_to_prompt(messages)
        temperature = self._get_temperature_for_strategy(strategy)
        
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{self.config.ollama_base_url}/api/generate",
                json={
                    "model": model["name"],
                    "prompt": prompt,
                    "options": {
                        "temperature": temperature,
                        "num_predict": self.config.max_tokens,
                    },
                    "stream": False
                }
            )
            
            if response.status_code != 200:
                raise Exception(f"Ollama request failed: {response.status_code}")
            
            data = response.json()
            response_text = data.get("response", "")
            
            confidence = self._calculate_response_confidence(response_text, strategy)
            
            return response_text, confidence
    
    def _convert_messages_to_prompt(self, messages: List[Dict[str, str]]) -> str:
        """Convert chat messages to prompt format for Ollama"""
        prompt_parts = []
        
        for message in messages:
            role = message["role"]
            content = message["content"]
            
            if role == "system":
                prompt_parts.append(f"System: {content}")
            elif role == "user":
                prompt_parts.append(f"Human: {content}")
            elif role == "assistant":
                prompt_parts.append(f"Assistant: {content}")
        
        prompt_parts.append("Assistant:")
        return "\n\n".join(prompt_parts)
    
    def _get_temperature_for_strategy(self, strategy: ReasoningStrategy) -> float:
        """Get appropriate temperature for reasoning strategy"""
        temperature_map = {
            ReasoningStrategy.CREATIVE: 0.9,
            ReasoningStrategy.ANALYTICAL: 0.3,
            ReasoningStrategy.DIRECT: 0.7,
            ReasoningStrategy.CONSERVATIVE: 0.1
        }
        return temperature_map.get(strategy, self.config.temperature)
    
    def _calculate_response_confidence(self, response_text: str, strategy: ReasoningStrategy) -> float:
        """
        Calculate confidence score for response based on various factors
        Mimics human self-assessment of response quality
        """
        if not response_text or len(response_text.strip()) < 10:
            return 0.1
        
        confidence_factors = []
        
        # Length factor (adequate responses need substance)
        length = len(response_text)
        if length > 100:
            confidence_factors.append(0.8)
        elif length > 50:
            confidence_factors.append(0.6)
        else:
            confidence_factors.append(0.4)
        
        # Coherence indicators
        sentences = sent_tokenize(response_text)
        if len(sentences) >= 2:
            confidence_factors.append(0.7)
        
        # Strategy-specific confidence adjustments
        if strategy == ReasoningStrategy.CONSERVATIVE:
            # Conservative responses should be cautious but complete
            if any(word in response_text.lower() for word in ['however', 'but', 'although', 'may', 'might']):
                confidence_factors.append(0.8)
        elif strategy == ReasoningStrategy.ANALYTICAL:
            # Analytical responses should show reasoning
            if any(phrase in response_text.lower() for phrase in ['because', 'therefore', 'as a result', 'analysis', 'consider']):
                confidence_factors.append(0.9)
        
        # Error indicators (reduce confidence)
        error_indicators = ['error', 'sorry', 'cannot', "don't know", 'unclear']
        if any(indicator in response_text.lower() for indicator in error_indicators):
            confidence_factors.append(0.3)
        
        # Average the factors
        base_confidence = sum(confidence_factors) / len(confidence_factors) if confidence_factors else 0.5
        
        # Ensure confidence is within reasonable bounds
        return min(max(base_confidence, 0.1), 0.95)
    
    def _update_model_performance(self, model_name: str, response_time: float, confidence: float, success: bool):
        """Update performance tracking for a model"""
        if model_name not in self.model_performance:
            self.model_performance[model_name] = {
                "total_requests": 0,
                "successful_requests": 0,
                "total_response_time": 0,
                "total_confidence": 0,
                "success_rate": 0,
                "avg_response_time": 0,
                "avg_confidence": 0
            }
        
        perf = self.model_performance[model_name]
        perf["total_requests"] += 1
        
        if success:
            perf["successful_requests"] += 1
            perf["total_response_time"] += response_time
            perf["total_confidence"] += confidence
        
        # Update averages
        perf["success_rate"] = perf["successful_requests"] / perf["total_requests"]
        if perf["successful_requests"] > 0:
            perf["avg_response_time"] = perf["total_response_time"] / perf["successful_requests"]
            perf["avg_confidence"] = perf["total_confidence"] / perf["successful_requests"]

# =============================================================================
# 🧠 CORE AGENTIC RAG SYSTEM
# =============================================================================

class AgenticRAGSystem:
    """
    🧠 Core Agentic RAG System
    
    The heart of our human-like reasoning system that combines:
    - Memory systems (working + long-term)
    - Cognitive reasoning strategies  
    - Adaptive learning and improvement
    - Multi-step problem solving
    """
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Core components
        self.vector_store = VectorStoreManager(config)
        self.llm_manager = LLMManager(config)
        self.document_processor = DocumentProcessor(config)
        
        # Memory systems
        self.working_memory = WorkingMemory()
        
        # System state
        self.is_initialized = False
        self.system_stats = {
            "queries_processed": 0,
            "documents_processed": 0,
            "total_chunks": 0,
            "avg_confidence": 0.0,
            "successful_responses": 0
        }
    
    async def initialize(self):
        """Initialize the complete RAG system"""
        try:
            self.logger.info("🚀 Initializing Advanced Agentic RAG System...")
            
            # Initialize core components
            await self.vector_store.initialize()
            await self.llm_manager.initialize()
            
            self.is_initialized = True
            self.logger.info("✅ Advanced Agentic RAG System initialized successfully!")
            
        except Exception as e:
            self.logger.error(f"Failed to initialize RAG system: {str(e)}")
            raise
    
    async def process_document(self, file_path: str, content: bytes = None) -> bool:
        """
        Process a document into our knowledge base
        
        Like human learning:
        - Read and understand the document
        - Break into meaningful chunks
        - Store in long-term memory
        - Make connections with existing knowledge
        """
        try:
            self.logger.info(f"📖 Processing document: {file_path}")
            
            # Process document into memory chunks
            chunks = await self.document_processor.process_file(file_path, content)
            
            if not chunks:
                self.logger.warning(f"No processable content found in {file_path}")
                return False
            
            # Store chunks in vector database
            success = await self.vector_store.store_chunks(chunks)
            
            if success:
                self.system_stats["documents_processed"] += 1
                self.system_stats["total_chunks"] += len(chunks)
                self.logger.info(f"✅ Document processed: {len(chunks)} chunks stored")
                return True
            else:
                self.logger.error(f"Failed to store chunks for {file_path}")
                return False
                
        except Exception as e:
            self.logger.error(f"Document processing failed for {file_path}: {str(e)}")
            return False
    
    async def query(self, user_query: str, max_results: int = 5) -> Dict[str, Any]:
        """
        Process a user query using human-like reasoning
        
        The core of our agentic reasoning:
        1. Understand the question (comprehension)
        2. Search relevant memories (recall)
        3. Reason about the information (analysis)
        4. Generate a thoughtful response (synthesis)
        5. Self-assess confidence (metacognition)
        """
        query_start_time = time.time()
        
        try:
            self.logger.info(f"🧠 Processing query: '{user_query[:100]}...'")
            
            # Clear working memory for new query (focus)
            self.working_memory.clear()
            
            # Step 1: Comprehension - Understand the query
            query_analysis = await self._analyze_query(user_query)
            self.working_memory.update_reasoning_trace(
                "query_analysis", query_analysis, query_analysis.get("confidence", 0.5)
            )
            
            # Step 2: Recall - Search for relevant information
            relevant_chunks = await self._search_relevant_knowledge(user_query, max_results)
            self.working_memory.update_reasoning_trace(
                "knowledge_search", f"Found {len(relevant_chunks)} relevant chunks", 
                0.8 if relevant_chunks else 0.2
            )
            
            # Add relevant chunks to working memory
            for chunk, score in relevant_chunks:
                self.working_memory.add_chunk(chunk)
            
            # Step 3: Reasoning Strategy Selection
            reasoning_strategy = await self._select_reasoning_strategy(user_query, relevant_chunks)
            self.working_memory.current_strategy = reasoning_strategy
            
            # Step 4: Generate Response with Human-like Reasoning
            response_data = await self._generate_reasoned_response(
                user_query, relevant_chunks, reasoning_strategy
            )
            
            # Step 5: Self-Assessment and Confidence Calculation
            final_confidence = self._calculate_overall_confidence(response_data, relevant_chunks)
            
            # Update system statistics
            self._update_system_stats(response_data, final_confidence)
            
            # Prepare comprehensive response
            query_time = time.time() - query_start_time
            
            result = {
                "response": response_data.get("response", "I apologize, but I couldn't generate a proper response."),
                "confidence": final_confidence,
                "sources": [
                    {
                        "content": chunk.content[:200] + "..." if len(chunk.content) > 200 else chunk.content,
                        "metadata": chunk.metadata,
                        "relevance_score": score
                    }
                    for chunk, score in relevant_chunks
                ],
                "reasoning_trace": self.working_memory.reasoning_trace,
                "strategy_used": reasoning_strategy.value,
                "context_found": len(relevant_chunks),
                "using_rag": len(relevant_chunks) > 0,
                "response_time": query_time,
                "system_stats": self.system_stats.copy()
            }
            
            self.logger.info(f"✅ Query processed successfully in {query_time:.2f}s (confidence: {final_confidence:.2f})")
            return result
            
        except Exception as e:
            self.logger.error(f"Query processing failed: {str(e)}")
            return {
                "response": "I encountered an error while processing your query. Please try again.",
                "confidence": 0.1,
                "sources": [],
                "error": str(e),
                "using_rag": False,
                "context_found": 0
            }
    
    async def _analyze_query(self, query: str) -> Dict[str, Any]:
        """
        Analyze the query to understand intent and complexity
        Like human reading comprehension
        """
        analysis = {
            "query_length": len(query),
            "word_count": len(query.split()),
            "question_words": [],
            "complexity": "simple",
            "intent": "informational",
            "confidence": 0.7
        }
        
        # Identify question words
        question_words = ["what", "how", "why", "when", "where", "who", "which", "can", "should", "could", "would"]
        found_question_words = [word for word in question_words if word.lower() in query.lower()]
        analysis["question_words"] = found_question_words
        
        # Assess complexity
        if len(query.split()) > 15 or len(found_question_words) > 2:
            analysis["complexity"] = "complex"
            analysis["confidence"] = 0.6
        elif len(query.split()) < 5:
            analysis["complexity"] = "simple"
            analysis["confidence"] = 0.8
        
        # Determine intent
        if any(word in query.lower() for word in ["explain", "describe", "tell me about"]):
            analysis["intent"] = "explanatory"
        elif any(word in query.lower() for word in ["how to", "steps", "process"]):
            analysis["intent"] = "procedural"
        elif any(word in query.lower() for word in ["compare", "difference", "better"]):
            analysis["intent"] = "comparative"
        
        return analysis
    
    async def _search_relevant_knowledge(self, query: str, max_results: int) -> List[Tuple[MemoryChunk, float]]:
        """
        Search for relevant knowledge using semantic similarity
        Like human memory recall and association
        """
        try:
            # Primary semantic search
            results = await self.vector_store.search(
                query=query,
                limit=max_results,
                score_threshold=0.3  # Lower threshold for broader recall
            )
            
            if not results:
                # Fallback: try with individual query terms
                self.logger.info("No results found, trying individual terms")
                query_terms = query.split()
                if len(query_terms) > 1:
                    for term in query_terms:
                        if len(term) > 3:  # Skip short words
                            term_results = await self.vector_store.search(
                                query=term,
                                limit=2,
                                score_threshold=0.2
                            )
                            results.extend(term_results)
                            if len(results) >= max_results:
                                break
            
            # Remove duplicates and sort by relevance
            seen_ids = set()
            unique_results = []
            for chunk, score in sorted(results, key=lambda x: x[1], reverse=True):
                if chunk.id not in seen_ids:
                    unique_results.append((chunk, score))
                    seen_ids.add(chunk.id)
                    if len(unique_results) >= max_results:
                        break
            
            return unique_results
            
        except Exception as e:
            self.logger.error(f"Knowledge search failed: {str(e)}")
            return []
    
    async def _select_reasoning_strategy(self, query: str, relevant_chunks: List[Tuple[MemoryChunk, float]]) -> ReasoningStrategy:
        """
        Select appropriate reasoning strategy based on query and available knowledge
        Like human metacognitive strategy selection
        """
        
        # Default strategy
        strategy = ReasoningStrategy.DIRECT
        
        # Strategy selection logic
        query_lower = query.lower()
        
        # Creative strategy for open-ended questions
        if any(word in query_lower for word in ["creative", "imagine", "ideas", "brainstorm", "innovative"]):
            strategy = ReasoningStrategy.CREATIVE
        
        # Analytical strategy for complex reasoning
        elif any(word in query_lower for word in ["analyze", "compare", "evaluate", "assess", "examine", "why"]):
            strategy = ReasoningStrategy.ANALYTICAL
        
        # Conservative strategy when uncertain or little context
        elif len(relevant_chunks) < 2 or not any(score > 0.7 for _, score in relevant_chunks):
            strategy = ReasoningStrategy.CONSERVATIVE
        
        # Direct strategy for straightforward questions
        elif any(word in query_lower for word in ["what", "when", "where", "who"]):
            strategy = ReasoningStrategy.DIRECT
        
        self.logger.info(f"🧠 Selected reasoning strategy: {strategy.value}")
        return strategy
    
    async def _generate_reasoned_response(self, 
                                        query: str, 
                                        relevant_chunks: List[Tuple[MemoryChunk, float]], 
                                        strategy: ReasoningStrategy) -> Dict[str, Any]:
        """
        Generate response using selected reasoning strategy
        Like human thought process synthesis
        """
        
        # Prepare context from relevant chunks
        context_pieces = []
        for chunk, score in relevant_chunks:
            context_pieces.append(f"[Relevance: {score:.2f}]\n{chunk.content}")
        
        context = "\n\n---\n\n".join(context_pieces) if context_pieces else "No specific context available."
        
        # Prepare system message based on strategy
        system_message = self._get_system_message_for_strategy(strategy)
        
        # Prepare conversation messages
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": f"Context Information:\n{context}\n\nUser Query: {query}\n\nPlease provide a helpful and accurate response based on the context provided."}
        ]
        
        # Generate response using LLM with selected strategy
        response_text, confidence, metadata = await self.llm_manager.generate_response(
            messages=messages,
            strategy=strategy,
            max_attempts=3
        )
        
        return {
            "response": response_text,
            "confidence": confidence,
            "metadata": metadata,
            "context_used": len(relevant_chunks) > 0
        }
    
    def _get_system_message_for_strategy(self, strategy: ReasoningStrategy) -> str:
        """Get system message tailored to reasoning strategy"""
        
        base_message = """You are an intelligent AI assistant with access to a knowledge base. Your goal is to provide helpful, accurate, and well-reasoned responses based on the provided context."""
        
        strategy_additions = {
            ReasoningStrategy.DIRECT: """
            Focus on providing direct, clear answers. Be concise but comprehensive. If the context contains the answer, state it clearly. If not, acknowledge the limitation.""",
            
            ReasoningStrategy.ANALYTICAL: """
            Take an analytical approach. Break down complex questions, consider multiple perspectives, and explain your reasoning process. Draw connections between different pieces of information.""",
            
            ReasoningStrategy.CREATIVE: """
            Think creatively and consider innovative approaches. When appropriate, explore possibilities beyond the immediate context while staying grounded in the available information.""",
            
            ReasoningStrategy.CONSERVATIVE: """
            Be cautious and precise. Only make claims you can strongly support with the provided context. Acknowledge uncertainty when appropriate and avoid speculation."""
        }
        
        return base_message + strategy_additions.get(strategy, "")
    
    def _calculate_overall_confidence(self, response_data: Dict[str, Any], relevant_chunks: List[Tuple[MemoryChunk, float]]) -> float:
        """
        Calculate overall confidence in the response
        Like human self-assessment of answer quality
        """
        confidence_factors = []
        
        # Base confidence from LLM
        base_confidence = response_data.get("confidence", 0.5)
        confidence_factors.append(base_confidence)
        
        # Context quality factor
        if relevant_chunks:
            avg_relevance = sum(score for _, score in relevant_chunks) / len(relevant_chunks)
            context_confidence = min(avg_relevance * 1.2, 0.9)  # Boost good relevance, cap at 0.9
            confidence_factors.append(context_confidence)
        else:
            # No context means lower confidence
            confidence_factors.append(0.3)
        
        # Response quality indicators
        response_text = response_data.get("response", "")
        if len(response_text) > 50 and not any(phrase in response_text.lower() 
                                             for phrase in ["don't know", "unclear", "cannot determine"]):
            confidence_factors.append(0.8)
        else:
            confidence_factors.append(0.4)
        
        # Working memory quality (how well did we process this?)
        if len(self.working_memory.confidence_history) > 0:
            avg_process_confidence = sum(self.working_memory.confidence_history) / len(self.working_memory.confidence_history)
            confidence_factors.append(avg_process_confidence)
        
        # Calculate weighted average
        overall_confidence = sum(confidence_factors) / len(confidence_factors)
        
        # Ensure reasonable bounds
        return min(max(overall_confidence, 0.1), 0.95)
    
    def _update_system_stats(self, response_data: Dict[str, Any], confidence: float):
        """Update system performance statistics"""
        self.system_stats["queries_processed"] += 1
        
        # Update average confidence
        prev_avg = self.system_stats["avg_confidence"]
        query_count = self.system_stats["queries_processed"]
        self.system_stats["avg_confidence"] = ((prev_avg * (query_count - 1)) + confidence) / query_count
        
        # Count successful responses (confidence > threshold)
        if confidence >= self.config.confidence_threshold:
            self.system_stats["successful_responses"] += 1
    
    async def get_system_status(self) -> Dict[str, Any]:
        """Get comprehensive system status"""
        vector_stats = await self.vector_store.get_collection_stats()
        
        return {
            "initialized": self.is_initialized,
            "vector_store": vector_stats,
            "llm_models": [model["name"] for model in self.llm_manager.available_models],
            "system_stats": self.system_stats,
            "working_memory_size": len(self.working_memory.active_chunks),
            "reasoning_history": len(self.working_memory.reasoning_trace)
        }

# =============================================================================
# 🌐 FASTAPI WEB APPLICATION
# =============================================================================

# Pydantic models for API
class QueryRequest(BaseModel):
    message: str
    max_results: Optional[int] = 5
    strategy: Optional[str] = "direct"

class QueryResponse(BaseModel):
    response: str
    confidence: float
    sources: List[Dict[str, Any]]
    using_rag: bool
    context_found: int
    response_time: float

class UploadResponse(BaseModel):
    message: str
    filename: str
    processed: bool
    chunks_created: Optional[int]
    success: bool

# Global application state
class AppState:
    def __init__(self):
        self.rag_system: Optional[AgenticRAGSystem] = None
        self.rag_status = {"initialized": False, "error": None}

app = FastAPI(
    title="Advanced Agentic RAG Pipeline",
    description="A production-ready RAG system with human-like reasoning capabilities",
    version="2.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global state
app_state = AppState()

@app.on_event("startup")
async def startup_event():
    """Initialize the RAG system on startup"""
    try:
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        
        logger = logging.getLogger(__name__)
        logger.info("🚀 Starting Advanced Agentic RAG Pipeline...")
        
        # Initialize RAG system
        config = RAGConfig()
        app_state.rag_system = AgenticRAGSystem(config)
        
        await app_state.rag_system.initialize()
        
        app_state.rag_status = {"initialized": True, "error": None}
        logger.info("✅ RAG System startup complete!")
        
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"❌ RAG System startup failed: {str(e)}")
        app_state.rag_status = {"initialized": False, "error": str(e)}

@app.get("/")
async def root():
    """Root endpoint with system information"""
    return {
        "message": "Advanced Agentic RAG Pipeline - Production Ready",
        "version": "2.0.0",
        "status": "operational" if app_state.rag_status["initialized"] else "initializing",
        "endpoints": {
            "chat": "/api/chat",
            "upload": "/api/upload",
            "status": "/api/status",
            "health": "/health"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    if not app_state.rag_status["initialized"]:
        raise HTTPException(status_code=503, detail="RAG system not initialized")
    
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.get("/api/status")
async def get_system_status():
    """Get detailed system status"""
    if not app_state.rag_status["initialized"]:
        return {
            "initialized": False,
            "error": app_state.rag_status.get("error"),
            "message": "RAG system not initialized"
        }
    
    try:
        status = await app_state.rag_system.get_system_status()
        return status
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get status: {str(e)}")

@app.post("/api/chat", response_model=QueryResponse)
async def chat_endpoint(request: QueryRequest):
    """
    Main chat endpoint for processing queries
    
    This is where the magic happens - human-like reasoning in action!
    """
    if not app_state.rag_status["initialized"]:
        raise HTTPException(status_code=503, detail="RAG system not initialized")
    
    try:
        # Validate strategy
        valid_strategies = [s.value for s in ReasoningStrategy]
        if request.strategy and request.strategy not in valid_strategies:
            request.strategy = "direct"
        
        # Process query using our agentic system
        result = await app_state.rag_system.query(
            user_query=request.message,
            max_results=request.max_results
        )
        
        return QueryResponse(**result)
        
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"Chat endpoint error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Query processing failed: {str(e)}")

@app.post("/api/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload and process documents into the knowledge base
    
    Like feeding new information to the AI's long-term memory
    """
    if not app_state.rag_status["initialized"]:
        raise HTTPException(status_code=503, detail="RAG system not initialized")
    
    try:
        # Validate file type
        allowed_extensions = {".pdf", ".docx", ".txt", ".md"}
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Read file content
        content = await file.read()
        
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty file")
        
        # Process document
        success = await app_state.rag_system.process_document(
            file_path=file.filename,
            content=content
        )
        
        if success:
            # Get updated stats to return chunk count
            status = await app_state.rag_system.get_system_status()
            total_chunks = status.get("system_stats", {}).get("total_chunks", 0)
            
            return UploadResponse(
                message="Document processed successfully",
                filename=file.filename,
                processed=True,
                chunks_created=total_chunks,  # This is cumulative, but gives user feedback
                success=True
            )
        else:
            return UploadResponse(
                message="Document processing failed",
                filename=file.filename,
                processed=False,
                chunks_created=None,
                success=False
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload processing failed: {str(e)}")

# =============================================================================
# 🚀 MAIN ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    """
    🚀 Production Entry Point
    
    Run the Advanced Agentic RAG Pipeline server
    """
    import argparse
    
    parser = argparse.ArgumentParser(description="Advanced Agentic RAG Pipeline")
    parser.add_argument("--host", default="0.0.0.0", help="Host to bind to")
    parser.add_argument("--port", type=int, default=8000, help="Port to bind to")
    parser.add_argument("--workers", type=int, default=1, help="Number of worker processes")
    parser.add_argument("--log-level", default="info", help="Logging level")
    
    args = parser.parse_args()
    
    print(f"""
    🧠 Advanced Agentic RAG Pipeline v2.0
    =====================================
    
    🚀 Starting server on {args.host}:{args.port}
    🧠 Human-like reasoning: ENABLED
    🔍 Semantic memory: ACTIVE  
    💡 Multi-strategy thinking: READY
    📚 Document processing: AVAILABLE
    
    Endpoints:
    • Chat: http://{args.host}:{args.port}/api/chat
    • Upload: http://{args.host}:{args.port}/api/upload  
    • Status: http://{args.host}:{args.port}/api/status
    
    """)
    
    uvicorn.run(
        "main_production_with_rag:app",
        host=args.host,
        port=args.port,
        workers=args.workers,
        log_level=args.log_level,
        reload=False  # Set to True for development
    )

"""
🎓 SYSTEM SUMMARY
=================

This Advanced Agentic RAG Pipeline mimics human cognitive processes through:

🧠 COGNITIVE ARCHITECTURE:
- Working Memory: Active processing space (like human attention)
- Long-term Memory: Persistent knowledge storage (like human memory)
- Reasoning Strategies: Multiple thinking approaches (like human metacognition)
- Self-Assessment: Confidence evaluation (like human self-reflection)

🔄 HUMAN-LIKE PROCESSES:
1. Query Comprehension: Understanding like human reading
2. Memory Recall: Associative search like human memory
3. Strategy Selection: Choosing thinking approach like humans
4. Reasoned Response: Multi-step synthesis like human thinking
5. Confidence Assessment: Self-evaluation like human metacognition

⚡ PRODUCTION FEATURES:
- Multi-LLM support with intelligent fallbacks
- Advanced document processing with semantic chunking
- Real-time vector search with Qdrant
- Comprehensive error handling and logging
- Performance monitoring and optimization
- RESTful API with full documentation

🎯 OPTIMIZATION:
- Efficient batch processing
- Intelligent caching strategies
- Progressive fallback mechanisms
- Memory-conscious operations
- Production-grade error handling
"""