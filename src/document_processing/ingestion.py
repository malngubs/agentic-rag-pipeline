"""
Document Processing Pipeline for RAG System.

This module handles intelligent document ingestion, chunking, and metadata extraction.
It processes multiple file formats and prepares them for semantic search.
"""

import asyncio
import logging
import hashlib
import time
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass
from enum import Enum

# File processing imports
try:
    import pymupdf4llm  # For PDF processing
except ImportError:
    print("Installing pymupdf4llm for PDF processing...")
    import subprocess
    subprocess.check_call(["pip", "install", "pymupdf4llm"])
    import pymupdf4llm

try:
    from docx import Document  # For DOCX processing
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document

import re
from datetime import datetime

# Configure logging
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    UNKNOWN = "unknown"


class ChunkStrategy(Enum):
    """Text chunking strategies."""
    FIXED_SIZE = "fixed_size"
    SEMANTIC = "semantic"
    PARAGRAPH = "paragraph"
    SENTENCE = "sentence"


@dataclass
class DocumentMetadata:
    """Document metadata extracted during processing."""
    filename: str
    file_type: DocumentType
    file_size: int
    created_at: datetime
    processed_at: datetime
    title: Optional[str] = None
    author: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: str = "en"
    content_hash: Optional[str] = None


@dataclass
class TextChunk:
    """A chunk of text with metadata."""
    id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None


@dataclass
class ProcessedDocument:
    """Complete processed document with chunks."""
    document_id: str
    metadata: DocumentMetadata
    chunks: List[TextChunk]
    processing_time: float
    success: bool
    error_message: Optional[str] = None


class DocumentProcessor:
    """
    Advanced document processor that handles multiple file formats
    and creates optimized chunks for RAG retrieval.
    """
    
    def __init__(self,
                 chunk_size: int = 512,
                 chunk_overlap: int = 64,
                 min_chunk_size: int = 50,
                 max_chunk_size: int = 2048):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size for text chunks
            chunk_overlap: Overlap between consecutive chunks
            min_chunk_size: Minimum acceptable chunk size
            max_chunk_size: Maximum acceptable chunk size
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        
        # Processing statistics
        self.stats = {
            "documents_processed": 0,
            "chunks_created": 0,
            "total_processing_time": 0.0,
            "file_type_counts": {},
            "error_count": 0
        }
    
    async def process_file(self, file_path: Union[str, Path]) -> ProcessedDocument:
        """
        Process a single file and return document chunks.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            ProcessedDocument with chunks and metadata
        """
        start_time = time.time()
        file_path = Path(file_path)
        
        try:
            # Validate file exists
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine file type
            file_type = self._detect_file_type(file_path)
            
            # Extract basic metadata
            metadata = self._extract_basic_metadata(file_path, file_type)
            
            # Extract text content
            text_content = await self._extract_text(file_path, file_type)
            
            # Update metadata with content analysis
            metadata.word_count = len(text_content.split())
            metadata.content_hash = hashlib.md5(text_content.encode()).hexdigest()
            
            # Create document ID
            document_id = f"doc_{metadata.content_hash[:12]}"
            
            # Create text chunks
            chunks = self._create_chunks(
                text_content, 
                document_id, 
                metadata
            )
            
            processing_time = time.time() - start_time
            
            # Update statistics
            self._update_stats(file_type, processing_time, len(chunks))
            
            logger.info(f"Processed {file_path.name}: {len(chunks)} chunks in {processing_time:.2f}s")
            
            return ProcessedDocument(
                document_id=document_id,
                metadata=metadata,
                chunks=chunks,
                processing_time=processing_time,
                success=True
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            self.stats["error_count"] += 1
            
            logger.error(f"Failed to process {file_path.name}: {e}")
            
            return ProcessedDocument(
                document_id=f"error_{int(time.time())}",
                metadata=DocumentMetadata(
                    filename=file_path.name,
                    file_type=DocumentType.UNKNOWN,
                    file_size=0,
                    created_at=datetime.now(),
                    processed_at=datetime.now()
                ),
                chunks=[],
                processing_time=processing_time,
                success=False,
                error_message=str(e)
            )
    
    def _detect_file_type(self, file_path: Path) -> DocumentType:
        """Detect file type from extension."""
        extension = file_path.suffix.lower()
        
        type_map = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,  # Treat .doc as .docx
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD,
            '.markdown': DocumentType.MD
        }
        
        return type_map.get(extension, DocumentType.UNKNOWN)
    
    def _extract_basic_metadata(self, file_path: Path, file_type: DocumentType) -> DocumentMetadata:
        """Extract basic file metadata."""
        stat = file_path.stat()
        
        return DocumentMetadata(
            filename=file_path.name,
            file_type=file_type,
            file_size=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            processed_at=datetime.now(),
            title=file_path.stem  # Use filename as default title
        )
    
    async def _extract_text(self, file_path: Path, file_type: DocumentType) -> str:
        """Extract text content from file based on type."""
        if file_type == DocumentType.PDF:
            return await self._extract_pdf_text(file_path)
        elif file_type == DocumentType.DOCX:
            return await self._extract_docx_text(file_path)
        elif file_type == DocumentType.TXT:
            return await self._extract_txt_text(file_path)
        elif file_type == DocumentType.MD:
            return await self._extract_markdown_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using pymupdf4llm."""
        try:
            # Use pymupdf4llm for better text extraction
            markdown_text = pymupdf4llm.to_markdown(str(file_path))
            
            # Clean up the markdown for better chunking
            text = self._clean_markdown_text(markdown_text)
            
            if not text.strip():
                raise ValueError("No text content found in PDF")
            
            return text
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            text = "\n\n".join(paragraphs)
            
            if not text.strip():
                raise ValueError("No text content found in DOCX")
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    if text.strip():
                        return text
                        
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise ValueError(f"Failed to extract text from TXT: {e}")
    
    async def _extract_markdown_text(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Clean up markdown formatting for better chunking
            text = self._clean_markdown_text(text)
            
            if not text.strip():
                raise ValueError("No text content found in Markdown")
            
            return text
            
        except Exception as e:
            logger.error(f"Markdown extraction failed: {e}")
            raise ValueError(f"Failed to extract text from Markdown: {e}")
    
    def _clean_markdown_text(self, text: str) -> str:
        """Clean markdown text for better processing."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Remove markdown headers but keep the text
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Remove markdown links but keep the text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Remove markdown emphasis but keep the text
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        return text.strip()
    
    def _create_chunks(self, 
                      text: str, 
                      document_id: str, 
                      metadata: DocumentMetadata) -> List[TextChunk]:
        """Create optimized text chunks from document text."""
        chunks = []
        
        # Split text into sentences for better chunking
        sentences = self._split_into_sentences(text)
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(potential_chunk) > self.chunk_size and current_chunk:
                # Create chunk from current content
                chunk = self._create_chunk(
                    current_chunk.strip(),
                    chunk_index,
                    current_start,
                    current_start + len(current_chunk),
                    document_id,
                    metadata
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                current_start = current_start + len(current_chunk) - len(overlap_text) - len(sentence)
                chunk_index += 1
            else:
                current_chunk = potential_chunk
        
        # Create final chunk if there's remaining content
        if current_chunk.strip() and len(current_chunk.strip()) >= self.min_chunk_size:
            chunk = self._create_chunk(
                current_chunk.strip(),
                chunk_index,
                current_start,
                current_start + len(current_chunk),
                document_id,
                metadata
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple heuristics."""
        # Simple sentence splitting - can be improved with more sophisticated NLP
        sentences = re.split(r'[.!?]+\s+', text)
        
        # Clean up sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 10:  # Minimum sentence length
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk."""
        if len(text) <= self.chunk_overlap:
            return text
        
        # Try to find a natural break point for overlap
        overlap_start = len(text) - self.chunk_overlap
        
        # Look for sentence boundary
        for i in range(overlap_start, len(text)):
            if text[i] in '.!?':
                return text[i+1:].strip()
        
        # Fallback to character-based overlap
        return text[-self.chunk_overlap:].strip()
    
    def _create_chunk(self, 
                     content: str,
                     chunk_index: int,
                     start_char: int,
                     end_char: int,
                     document_id: str,
                     metadata: DocumentMetadata) -> TextChunk:
        """Create a TextChunk object."""
        # Generate proper UUID for chunk ID
        chunk_id = str(uuid.uuid4())
        
        chunk_metadata = {
            "document_id": document_id,
            "filename": metadata.filename,
            "file_type": metadata.file_type.value,
            "chunk_index": chunk_index,
            "file_size": metadata.file_size,
            "created_at": metadata.created_at.isoformat(),
            "processed_at": metadata.processed_at.isoformat(),
            "title": metadata.title,
            "author": metadata.author,
            "language": metadata.language,
            "word_count": len(content.split()),
            "char_count": len(content)
        }
        
        return TextChunk(
            id=chunk_id,
            content=content,
            chunk_index=chunk_index,
            start_char=start_char,
            end_char=end_char,
            metadata=chunk_metadata
        )
    
    def _update_stats(self, file_type: DocumentType, processing_time: float, chunk_count: int):
        """Update processing statistics."""
        self.stats["documents_processed"] += 1
        self.stats["chunks_created"] += chunk_count
        self.stats["total_processing_time"] += processing_time
        
        type_key = file_type.value
        if type_key not in self.stats["file_type_counts"]:
            self.stats["file_type_counts"][type_key] = 0
        self.stats["file_type_counts"][type_key] += 1
    
    async def process_directory(self, directory_path: Union[str, Path]) -> List[ProcessedDocument]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to directory containing files
            
        Returns:
            List of ProcessedDocument objects
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Directory not found: {directory_path}")
        
        # Supported file extensions
        supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.markdown'}
        
        # Find all supported files
        files_to_process = []
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                files_to_process.append(file_path)
        
        logger.info(f"Found {len(files_to_process)} files to process in {directory_path}")
        
        # Process files
        results = []
        for file_path in files_to_process:
            try:
                result = await self.process_file(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                # Continue processing other files
        
        return results
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics."""
        stats = self.stats.copy()
        
        if stats["documents_processed"] > 0:
            stats["avg_processing_time"] = stats["total_processing_time"] / stats["documents_processed"]
            stats["avg_chunks_per_document"] = stats["chunks_created"] / stats["documents_processed"]
        else:
            stats["avg_processing_time"] = 0.0
            stats["avg_chunks_per_document"] = 0.0
        
        return stats


# Convenience functions
async def process_single_file(file_path: Union[str, Path], 
                            chunk_size: int = 512,
                            chunk_overlap: int = 64) -> ProcessedDocument:
    """Process a single file with default settings."""
    processor = DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return await processor.process_file(file_path)


# Example usage and testing
if __name__ == "__main__":
    async def test_document_processor():
        """Test the document processor with sample files."""
        print("Testing Document Processor...")
        
        processor = DocumentProcessor(chunk_size=256, chunk_overlap=32)
        
        # Create test directory and sample files
        test_dir = Path("./test_documents")
        test_dir.mkdir(exist_ok=True)
        
        # Create sample text file
        sample_txt = test_dir / "sample.txt"
        with open(sample_txt, 'w') as f:
            f.write("""
            This is a sample document for testing the document processor.
            
            It contains multiple paragraphs with different content.
            The processor should split this into appropriate chunks.
            
            Each chunk should maintain context while staying within size limits.
            This helps ensure effective retrieval during RAG operations.
            
            The system can handle various file formats including PDF, DOCX, and plain text.
            """)
        
        # Process the file
        result = await processor.process_file(sample_txt)
        
        print(f"Processing result: {result.success}")
        print(f"Document ID: {result.document_id}")
        print(f"Processing time: {result.processing_time:.3f}s")
        print(f"Number of chunks: {len(result.chunks)}")
        
        for i, chunk in enumerate(result.chunks):
            print(f"\nChunk {i}:")
            print(f"  ID: {chunk.id}")
            print(f"  Content: {chunk.content[:100]}...")
            print(f"  Word count: {chunk.metadata['word_count']}")
        
        # Get statistics
        stats = processor.get_processing_stats()
        print(f"\nProcessing stats: {stats}")
    
    # Run test
    asyncio.run(test_document_processor())